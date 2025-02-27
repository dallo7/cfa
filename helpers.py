import time
from docx import Document
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from openai import OpenAI
import os
from dotenv import load_dotenv

load_dotenv()

# YOUR_API_KEY = os.getenv("YOUR_API_KEY")
YOUR_API_KEY="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"

def determine_ai_toolkit_type(user_input):
    messages = [
        {
            "role": "system",
            "content": (
                "You are an AI toolkit advisor, specializing in creating AI solutions for different domains. Your task is to determine the type of AI toolkit a user wants to build based on their input. The final output should categorize their needs and suggest relevant AI applications. Ensure the language follows UK English conventions and that the content is structured, actionable, and engaging."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {user_input}\n\n"
                "Task: Analyse the user's input and determine the appropriate AI toolkit category. Provide recommendations in the following structure:\n\n"
                "1. Identified AI Toolkit Category\n"
                "   - Classify the toolkit type (e.g., AI for Journalism, AI for Business Analytics, AI for Healthcare, etc.).\n\n"
                "2. Core AI Technologies Required\n"
                "   - List essential AI technologies such as Machine Learning, Natural Language Processing, Computer Vision, etc.\n\n"
                "3. Relevant AI Use Cases\n"
                "   - Describe key applications for the selected domain (e.g., AI-powered fact-checking for journalism, predictive analytics for business, diagnostic support for healthcare).\n\n"
                "4. Recommended Tools and Platforms\n"
                "   - Provide an overview of suitable AI tools, including open-source and commercial options.\n\n"
                "5. Implementation Considerations\n"
                "   - Highlight challenges, best practices, and ethical considerations.\n\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
                "Output: An 8 sentence summary of the above points"
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")

    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data



def introduction_to_ai_and_context(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write an introduction to AI and its context in journalism, "
                "specifically tailored for African journalists. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write an introduction to AI and its context in journalism, specifically tailored for African journalists using the following guidelines:\n\n"
                "1.  Key AI Concepts\n"
                "    -   Core AI Technologies: Artificial Intelligence, Machine Learning, Generative AI, Natural Language Processing (NLP).\n"
                "    -   Discuss the above terms using atleast 3 lines each and make it easy to understand for a journalist"
                "    -   For each term provide atleast 2 citations/urls"
                "    -   AI's Role in Journalism: How AI transforms tasks, improves accuracy, and drives efficiency.\n"
                "        * Provide African examples and context (e.g., use of AI for combating fake news during elections, automating transcription of local dialects)."
                "2.  Domain-Specific Concepts\n"
                "    -   Relevant Jargon and Terms: Introduce specialised terminology, explain their meanings, and show how they relate to AI tools.\n"
                "    -   Examples and Use Cases: Provide context for how these terms apply in real-world scenarios.\n"
                "        * Focus on challenges and opportunities specific to African communities and journalists (e.g., addressing misinformation in low-resource settings, promoting local languages)."
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ai_applications_in_journalism(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on AI applications in journalism, "
                "with a focus on African context and examples. "
                "The final document must be in UK English, not exceed 4 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on AI applications in journalism using the following guidelines:\n\n"
                "1.  Use Cases and Strategies\n"
                "    -   Practical Applications: Showcase practical AI applications in journalism (e.g., fact-checking, audience segmentation, content generation).\n"
                "        * Provide African examples (e.g., using AI to track wildlife poaching, analyze social media trends during elections, generate news summaries in local languages)."
                "    -   AI's Potential: Discuss AI’s potential for enhancing productivity, innovation, and outcomes in African journalism.\n"
                "2.  Tools and Platforms\n"
                "    -   Overview of Tools: Provide an overview of relevant AI tools for journalists (cloud-based, open-source, or self-hosted).\n"
                "        * Include at least 32 AI tools across 8 categories (e.g., content generation, fact-checking, translation) that are relevant to African journalists.\n"
                "        * Prioritize tools that are easy to use(provide a description), open-source, or cost-effective.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ai_readiness_and_risk_assessment(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on AI readiness and risk assessment for newsrooms, "
                "with an African perspective. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on AI readiness and risk assessment using the following guidelines:\n\n"
                "1.  Adoption Framework\n"
                "    -   Readiness Steps: Provide at least 3 AI readiness tests/guides to assess organisational readiness for AI implementation in African newsrooms.\n"
                "    -   Risk Identification: Provide at least 3 risk assessment tests/guides to identify potential pitfalls and ethical concerns specific to the African context (e.g., data bias, cultural sensitivity, limited resources).\n"
                "2.  Evaluation Metrics\n"
                "    -   Metrics for Success: Define metrics for evaluating the success of AI adoption in African journalism (e.g., increased efficiency, improved accuracy, enhanced reach to underserved communities).\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def practical_guidance_for_ai_use(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write practical guidance on using AI tools in journalism, "
                "including prompt engineering and operational considerations, with an African perspective. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write practical guidance on AI use in journalism using the following guidelines:\n\n"
                "1.  Prompt Engineering\n"
                "    -   Effective Prompts: Offer guidance on crafting effective prompts for large language models (LLMs) and AI systems, including templates and engaging examples relevant to African journalism.\n"
                "2.  Operational Considerations\n"
                "    -   Workflow Integration: Explain how to integrate AI tools into existing workflows in African newsrooms.\n"
                "    -   Challenges: Address challenges such as AI bias, data quality, and scalability in the African context.\n"
                "3.  Examples of Resources and Databases that can help Journalists to prompt\n"
                "    -   Opensource materials for journalist\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def ethical_and_responsible_ai_use(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on ethical and responsible AI use in journalism, "
                "with a focus on African considerations. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on ethical and responsible AI use in journalism using the following guidelines:\n\n"
                "1.  Ethical Principles\n"
                "    -   Core Values: Discuss fairness, transparency, and accountability in AI deployment within the African context.\n"
                "    -   Critical Issues: Address source verification, misinformation, and bias mitigation specific to African communities and media landscapes.\n"
                "2.  Regulatory Considerations\n"
                "    -   Compliance: Highlight compliance with relevant standards and legal frameworks, including those emerging in African countries.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def challenges_and_future_trends(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on challenges and future trends of AI in journalism, "
                "with an African perspective. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on challenges and future trends of AI in journalism using the following guidelines:\n\n"
                "1.  Current Challenges\n"
                "    -   Limitations: Discuss limitations of AI tools in journalism (e.g., accuracy, scalability, over-reliance).\n"
                "        * Focus on challenges specific to African media (e.g., lack of local language support, limited data availability, digital divide).\n"
                "2.  Emerging Trends\n"
                "    -   Future Advancements: Explore future possibilities, such as advancements in AI models or new applications relevant to African journalism.\n"
                "3.  Case Studies\n"
                "    -   Examples: Include case studies showing both successful and failed implementations of AI in African journalism (e.g., initiatives to combat misinformation, projects promoting local languages).\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


def resources_and_training_materials(query):
    messages = [
        {
            "role": "system",
            "content": (
                "You are a chief researcher and an expert in building AI-Focused Toolkits that help journalists improve productivity "
                "and integrate open-source tools into their workflows. Your task is to write a section on resources and training materials for journalists interested in AI, "
                "with an African focus. "
                "The final document must be in UK English, not exceed 3 pages, use Times Roman font, "
                "and cite sources via clickable links rather than APA referencing. Tailor the content to be engaging and actionable for a journalistic audience."
            ),
        },
        {
            "role": "user",
            "content": (
                f"Input: {query}\n\n"
                "Task: Write a section on resources and training materials using the following guidelines:\n\n"
                "1.  Further Learning\n"
                "    -   Provide 10 curated further learning materials (links to research papers, articles, courses, etc.) relevant to AI in journalism, with some specifically focused on the African context.\n"
                "2.  Exercises and Workshops\n"
                "    -   Include 5 hands-on exercises or workshop guides for practical learning about AI tools and techniques, tailored for African journalists.\n"
                "3.  Case Studies\n"
                "    -   Present 5 real-world case studies illustrating effective AI deployment in journalism, with examples from African newsrooms or initiatives.\n"
                "THE OUTPUT SHOULD BE CLEAN TEXT WITHOUT '#' "
            ),
        }
    ]

    client = OpenAI(api_key=YOUR_API_KEY, base_url="https://api.perplexity.ai")
    response = client.chat.completions.create(
        model="sonar-pro",
        messages=messages,
        # timeout=60
    )

    extracted_data = response.choices[0].message.content

    return extracted_data


# Creat & Style docx!

def set_document_styles(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)


def add_table_of_contents(doc):
    doc.add_paragraph("Table of Contents", style='Heading 1')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    doc.add_page_break()


def add_section(doc, title, content, level=1):
    doc.add_heading(title, level=level)
    doc.add_paragraph(content)
    doc.add_paragraph()

if __name__ == "__main__":
    pass
