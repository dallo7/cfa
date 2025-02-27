import dash
import dash_bootstrap_components as dbc
from dash import dcc, html, Input, Output, State
import time
import base64
import multiprocessing
import io
import dash_auth
import Dashauth
from docx import Document
from helpers import *

YOUR_API_KEY="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"

# This function processes one section only!
def process_section(args):
    title, function, optimisedQuery = args
    content = function(optimisedQuery)
    return (title, content)

# --- Document Generation Function ---
def generate_toolkit_document(query):
    optimisedQuery = determine_ai_toolkit_type(query)
    sections = [
        ("Introduction to AI and Context", introduction_to_ai_and_context),
        ("AI Applications in Journalism", ai_applications_in_journalism),
        ("AI Readiness and Risk Assessment", ai_readiness_and_risk_assessment),
        ("Practical Guidance for AI Use", practical_guidance_for_ai_use),
        ("Ethical and Responsible AI Use", ethical_and_responsible_ai_use),
        ("Challenges and Future Trends", challenges_and_future_trends),
        ("Resources and Training Materials", resources_and_training_materials),
    ]
    doc = Document()
    set_document_styles(doc)
    doc.add_heading("CFA SANDBOX-GENERATED TOOLKIT", level=0)
    add_table_of_contents(doc)

    pool = multiprocessing.Pool(processes=len(sections))
    args_list = [(title, func, optimisedQuery) for title, func in sections]
    results = pool.map(process_section, args_list)
    pool.close()
    pool.join()

    for title, content in results:
        add_section(doc, title, content, level=1)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# --- Dash App Setup ---
app = dash.Dash(__name__, external_stylesheets=[dbc.themes.BOOTSTRAP])

auth = dash_auth.BasicAuth(
    app,
    Dashauth.VALID_USERNAME_PASSWORD_PAIRS
)

server = app.server

app.layout = html.Div([
    html.H1("CFA SANDBOX AUT0-TOOLKIT POC", style={
        'textAlign': 'center',
        'color': '#333366',
        'marginBottom': '20px'
    }),
    html.Div([
        html.Label("Enter your prompt:", style={'fontWeight': 'bold', 'marginBottom': '10px'}),
        dcc.Textarea(
            id='prompt-input',
            rows=5,
            style={'width': '100%', 'borderRadius': '10px', 'padding': '10px'}
        ),
        html.Br(),
        html.Button(
            'Generate Toolkit',
            id='generate-button',
            n_clicks=0,
            style={
                'backgroundColor': '#4CAF50',
                'color': 'white',
                'padding': '10px 20px',
                'border': 'none',
                'borderRadius': '10px',
                'cursor': 'pointer',
                'marginTop': '10px'
            }
        )
    ], style={'margin': '20px'}),

    dcc.Loading(
        id="loading-1",
        type="default",
        children=[
            html.Div(
                [
                    "👋🏿 Thinking! and Retrieving data for generating the Toolkit will take less than 3 mins 😇",
                    html.Div(
                        " Sit Tight",
                        style={
                            'textAlign': 'center',
                            'margin': '10px',
                            'fontSize': '18px',
                            'color': '#fff',
                            'backgroundColor': '#28a745',
                            'padding': '10px',
                            'borderRadius': '5px',
                            'display': 'inline-block',
                            'cursor': 'pointer'
                        }
                    )
                ],
                style={'textAlign': 'center', 'margin': '10px', 'fontSize': '18px', 'color': '#333'}
            ),
            html.Div(id='progress-container', style={'margin': '20px'}),
            html.Div(id='chain-of-thought-output', style={
                'whiteSpace': 'pre-line',
                'margin': '20px',
                'backgroundColor': '#f4f4f4',
                'padding': '15px',
                'borderRadius': '10px'
            }),
            html.Div(id='download-link-container', style={'margin': '20px'})
        ]
    )
], style={
    'fontFamily': 'Arial',
    'width': '75%',
    'margin': '50px auto',
    'padding': '20px',
    'backgroundColor': '#fff',
    'borderRadius': '10px',
    'boxShadow': '0 0 10px rgba(0,0,0,0.1)'
})


@app.callback(
    [Output('progress-container', 'children'),
     Output('chain-of-thought-output', 'children'),
     Output('download-link-container', 'children'),
     Output('generate-button', 'disabled')],
    [Input('generate-button', 'n_clicks')],
    [State('prompt-input', 'value')]
)
def update_output(n_clicks, prompt):
    if n_clicks == 0 or prompt is None:
        return "", "", "", False

    progress_updates = []
    thought_process = []

    # Initial Progress
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=0, max=100, style={'width': '100%'}),
        html.P("Initializing process...")
    ]))
    thought_process.append("1. Optimizing your query for the AI model...\n")
    time.sleep(2)

    # Stage 1: Optimizing query
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=10, max=100, style={'width': '100%'}),
        html.P("Optimizing query...")
    ]))

    # Simulate section generation (stages 2-8)
    section_names = [
        "Introduction to AI and Context", "AI Applications in Journalism",
        "AI Readiness and Risk Assessment", "Practical Guidance for AI Use",
        "Ethical and Responsible AI Use", "Challenges and Future Trends",
        "Resources and Training Materials"
    ]
    for i, section_name in enumerate(section_names):
        thought_process.append(f"{i + 2}. Generating section: {section_name}...\n")
        time.sleep(2)
        progress_value = 10 + (i + 1) * (70 / len(section_names))
        progress_updates.append(html.Div([
            dbc.Progress(id='progress-bar', value=progress_value, max=100, style={'width': '100%'}),
            html.P(f"Generating {section_name}...")
        ]))

    # Stage 9: Assembling document
    thought_process.append("9. Assembling the document...\n")
    time.sleep(2)
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=90, max=100, style={'width': '100%'}),
        html.P("Finalizing document...")
    ]))

    # Stage 10: Preparing document for download
    thought_process.append("10. Preparing document for download...\n")
    buffer = generate_toolkit_document(prompt)
    progress_updates.append(html.Div([
        dbc.Progress(id='progress-bar', value=100, max=100, style={'width': '100%'}),
        html.P("Document ready!")
    ]))

    # Create a base64 encoded string of the document
    buffer.seek(0)
    doc_bytes = buffer.read()
    b64 = base64.b64encode(doc_bytes).decode()
    download_link = html.A(
        'Download Your Toolkit',
        id='download-link',
        download="AI4Democracy_Toolkit.docx",
        href=f"data:application/vnd.openxmlformats-officedocument.wordprocessingml.document;base64,{b64}",
        target="_blank",
        style={
            'backgroundColor': '#007BFF',
            'color': 'white',
            'padding': '10px 20px',
            'borderRadius': '5px',
            'textDecoration': 'none',
            'display': 'inline-block'
        }
    )

    return progress_updates, "".join(thought_process), download_link, False


if __name__ == '__main__':
    app.run_server(debug=True, port=8000)
